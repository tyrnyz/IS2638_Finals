import os
import tempfile
import shutil
import time
import csv
import traceback
from typing import List, Dict, Any, Optional
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Request
from fastapi.responses import JSONResponse, PlainTextResponse
from fastapi.exceptions import RequestValidationError
from starlette.exceptions import HTTPException as StarletteHTTPException
from dotenv import load_dotenv
from datetime import datetime, timezone
from io import BytesIO, StringIO

# DOCX helper
from docx import Document

# Load .env
load_dotenv()

# ETL modules and supabase client (your project modules)
from backend.app.etl import (
    airlines_etl,
    passengers_etl,
    flights_etl,
    airports_etl,
    travelagency_etl,
    corporatesales_etl,
)
from backend.app.services.supabase_client import sb

# FastAPI + CORS
from fastapi.middleware.cors import CORSMiddleware

app = FastAPI(title="ETL Upload API (IS2638) - staged upload + explicit process")

FRONTEND_URL = os.getenv("FRONTEND_URL", "http://localhost:5173")
app.add_middleware(
    CORSMiddleware,
    allow_origins=[FRONTEND_URL],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# -----------------------
# Global exception handlers (return JSON consistently)
# -----------------------
@app.exception_handler(StarletteHTTPException)
async def http_exception_handler(request: Request, exc: StarletteHTTPException):
    # Return JSON instead of HTML for 404/405/etc
    return JSONResponse(
        status_code=exc.status_code,
        content={"status": "error", "message": exc.detail if exc.detail else "HTTP error"},
    )

@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request: Request, exc: RequestValidationError):
    return JSONResponse(
        status_code=422,
        content={"status": "error", "message": "Validation error", "details": exc.errors()},
    )

@app.exception_handler(Exception)
async def general_exception_handler(request: Request, exc: Exception):
    # log server-side traceback for debugging
    traceback.print_exc()
    return JSONResponse(
        status_code=500,
        content={"status": "error", "message": "internal server error", "detail": str(exc)},
    )

# -----------------------
# Quick debug endpoints
# -----------------------
@app.get("/")
async def root():
    return JSONResponse({"status": "ok", "app": app.title})

@app.get("/api/routes")
async def list_routes():
    routes = []
    for r in app.routes:
        # route.path_format for FastAPI Starlette Route; fallback to r.path
        try:
            path = getattr(r, "path", None) or getattr(r, "path_format", None) or str(r)
        except Exception:
            path = str(r)
        methods = []
        try:
            methods = list(r.methods) if getattr(r, "methods", None) else []
        except Exception:
            methods = []
        routes.append({"path": path, "methods": methods, "name": getattr(r, "name", "")})
    return JSONResponse({"status": "ok", "routes": routes})

# dataset -> cleaned table, rpc, module (single-run rpc)
DATASET_MAP = {
    "airline": {
        "cleaned_table": "cleaned_airlines",
        "rpc": "process_cleaned_airlines",
        "etl_module": airlines_etl,
    },
    "airlines": {
        "cleaned_table": "cleaned_airlines",
        "rpc": "process_cleaned_airlines",
        "etl_module": airlines_etl,
    },
    "passenger": {
        "cleaned_table": "cleaned_passengers",
        "rpc": "process_cleaned_passengers",
        "etl_module": passengers_etl,
    },
    "passengers": {
        "cleaned_table": "cleaned_passengers",
        "rpc": "process_cleaned_passengers",
        "etl_module": passengers_etl,
    },
    "flight": {
        "cleaned_table": "cleaned_flights",
        "rpc": "process_cleaned_flights",
        "etl_module": flights_etl,
    },
    "flights": {
        "cleaned_table": "cleaned_flights",
        "rpc": "process_cleaned_flights",
        "etl_module": flights_etl,
    },
    "airport": {
        "cleaned_table": "cleaned_airports",
        "rpc": "process_cleaned_airports",
        "etl_module": airports_etl,
    },
    "airports": {
        "cleaned_table": "cleaned_airports",
        "rpc": "process_cleaned_airports",
        "etl_module": airports_etl,
    },
    "travelagency": {
        "cleaned_table": "cleaned_travelagency",
        "rpc": "process_cleaned_travelagency",
        "etl_module": travelagency_etl,
    },
    "travel_agency": {
        "cleaned_table": "cleaned_travelagency",
        "rpc": "process_cleaned_travelagency",
        "etl_module": travelagency_etl,
    },
    "corporatesales": {
        "cleaned_table": "cleaned_corporatesales",
        "rpc": "process_cleaned_corporatesales",
        "etl_module": corporatesales_etl,
    },
    "corporate_sales": {
        "cleaned_table": "cleaned_corporatesales",
        "rpc": "process_cleaned_corporatesales",
        "etl_module": corporatesales_etl,
    },
}

# config (env overrides)
BATCH_INSERT_SIZE = int(os.getenv("BATCH_INSERT_SIZE", "200"))
MAX_FILE_BYTES = int(os.getenv("MAX_FILE_BYTES", str(10 * 1024 * 1024)))  # default 10MB
STORE_UPLOADS = os.getenv("STORE_UPLOADS", "false").lower() in ("1", "true", "yes")

# -----------------------
# Helpers
# -----------------------
def batch_insert(table_name: str, records: List[Dict[str, Any]], batch_size: int = BATCH_INSERT_SIZE) -> int:
    if not records:
        return 0
    inserted = 0
    for i in range(0, len(records), batch_size):
        chunk = records[i : i + batch_size]
        res = sb.table(table_name).insert(chunk).execute()
        if isinstance(res, dict) and res.get("error"):
            raise RuntimeError(f"Error inserting into {table_name}: {res.get('error')}")
        if hasattr(res, "error") and res.error:
            raise RuntimeError(f"Error inserting into {table_name}: {res.error}")
        inserted += len(chunk)
    return inserted


def parse_rpc_count(res) -> int:
    try:
        if hasattr(res, "data"):
            d = res.data
            if isinstance(d, list) and d:
                first = d[0]
                if isinstance(first, dict):
                    for v in first.values():
                        if isinstance(v, int):
                            return v
                if isinstance(first, int):
                    return first
            if isinstance(d, int):
                return d
    except Exception:
        pass
    if isinstance(res, dict):
        if "data" in res:
            d = res["data"]
            if isinstance(d, list) and d:
                first = d[0]
                if isinstance(first, dict):
                    for v in first.values():
                        if isinstance(v, int):
                            return v
                if isinstance(first, int):
                    return first
            if isinstance(d, int):
                return d
        for k in ("count", "rows_affected", "result"):
            if k in res and isinstance(res[k], int):
                return res[k]
    return 0

# -----------------------
# Validation: required fields per dataset (upload-time validation)
# -----------------------
# Each inner list is an OR-group. At least one field from each inner list must be present/non-empty.
# NOTE: alliance was removed from the airline REQUIRED_FIELDS as requested (alliance is optional).
REQUIRED_FIELDS = {
    "airline": [["airlinekey"], ["airlinename"]],
    "airlines": [["airlinekey"], ["airlinename"]],
    "passenger": [["passengerkey"], ["fullname"]],
    "passengers": [["passengerkey"], ["fullname"]],
    "flight": [["flightkey"], ["originairportkey", "origin_airportkey", "origin"], ["destinationairportkey", "destination_airportkey", "destination"]],
    "flights": [["flightkey"], ["originairportkey", "origin"], ["destinationairportkey", "destination"]],
    "airport": [["airportkey"], ["airportname"], ["city", "country"]],
    "airports": [["airportkey"], ["airportname"], ["city", "country"]],
    "travelagency": [["agency"], ["saleamount", "sale_amount"], ["saledate", "sale_date"]],
    "travel_agency": [["agency"], ["saleamount", "sale_amount"], ["saledate", "sale_date"]],
    "corporatesales": [["invoice"], ["transactionid"], ["saleamount", "sale_amount", "saledate", "sale_date"]],
    "corporate_sales": [["invoice"], ["transactionid"], ["saleamount", "sale_amount", "saledate", "sale_date"]],
}

def validate_required_fields(dataset_key: str, row: Dict[str, Any]) -> Optional[str]:
    """
    Return None if row is valid, otherwise return a short error message describing missing required fields.
    Row keys are expected to be normalized (lowercase, underscores) by parse_csv_text_to_dicts.
    """
    dataset_key = (dataset_key or "").lower()
    groups = REQUIRED_FIELDS.get(dataset_key)
    if not groups:
        return None  # no validation rules for this dataset

    missing_groups = []
    for group in groups:
        satisfied = False
        for key in group:
            v = row.get(key)
            if v is not None and str(v).strip() != "":
                satisfied = True
                break
        if not satisfied:
            missing_groups.append(group)

    if not missing_groups:
        return None

    readable = []
    for g in missing_groups:
        readable.append("(" + " OR ".join(g) + ")")
    return "missing required fields: " + ", ".join(readable)

def call_rpc_once(rpc_name: str, p_upload_id: Optional[int] = None) -> int:
    if p_upload_id is None:
        res = sb.rpc(rpc_name, {"p_upload_id": None}).execute()
    else:
        res = sb.rpc(rpc_name, {"p_upload_id": p_upload_id}).execute()
    return parse_rpc_count(res)


def insert_etl_run(jobname: str, status: str, note: str = None) -> int:
    try:
        rec = {
            "jobname": jobname,
            "status": status,
            "note": note,
        }
        res = sb.table("etl_runs").insert(rec).execute()
        if isinstance(res, dict) and res.get("error"):
            print("Warning: could not insert etl_runs:", res.get("error"))
            return -1
        if hasattr(res, "data") and isinstance(res.data, list) and res.data:
            return res.data[0].get("id", -1)
    except Exception as e:
        print("Warning: insert_etl_run failed:", e)
    return -1


def safe_update_etl_run(run_id: int, status: str, note: Optional[str] = None) -> bool:
    ts = datetime.now(timezone.utc).isoformat()
    tried = []
    for finished_col in ("finishedat", "finished_at", "finished"):
        payload = {"status": status, finished_col: ts}
        if note is not None:
            payload["note"] = str(note)
        try:
            sb.table("etl_runs").update(payload).eq("id", run_id).execute()
            return True
        except Exception as e:
            tried.append((finished_col, str(e)))
            continue
    try:
        payload = {"status": status}
        if note is not None:
            payload["note"] = str(note)
        sb.table("etl_runs").update(payload).eq("id", run_id).execute()
        return True
    except Exception as e:
        print("Warning: could not update etl_runs. Attempts:", tried, "final:", str(e))
        return False


# -----------------------
# DOCX -> CSV converter (table first; paragraph fallback)
# -----------------------
MAX_DOCX_SIZE = 10 * 1024 * 1024  # 10MB default

ALLOWED_DOCX_CONTENT_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/octet-stream",
}


def escape_csv_cell(cell_text: str) -> str:
    if cell_text is None:
        cell_text = ""
    text = str(cell_text).replace("\r", " ").replace("\n", " ").strip()
    if '"' in text:
        text = text.replace('"', '""')
    if ("," in text) or ('"' in text) or ("\n" in text) or ("\r" in text):
        return f'"{text}"'
    return text


def docx_to_csv_text_with_fallback(fileobj: BytesIO, table_selection: str = "first") -> str:
    doc = Document(fileobj)

    if doc.tables:
        csv_lines = []
        tables = doc.tables if table_selection == "all" else [doc.tables[0]]
        for t in tables:
            for r in t.rows:
                cells = [c.text for c in r.cells]
                escaped = [escape_csv_cell(c) for c in cells]
                csv_lines.append(",".join(escaped))
            if table_selection == "all":
                csv_lines.append("")
        return "\n".join(csv_lines)

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    if not paragraphs:
        raise ValueError("No tables found in DOCX and no readable paragraph text to convert to CSV.")

    text_blob = "\n".join(paragraphs)
    if any(sep in text_blob for sep in [",", ";", "\t"]):
        return text_blob

    rows = []
    for p in paragraphs:
        cols = p.split()
        escaped = [escape_csv_cell(c) for c in cols]
        rows.append(",".join(escaped))
    return "\n".join(rows)


# -----------------------
# Utility: parse CSV text into list[dict]
# -----------------------
def parse_csv_text_to_dicts(csv_text: str) -> List[Dict[str, Any]]:
    """
    Uses csv.DictReader to parse CSV text into list of dicts.
    Trims keys and values and normalizes empty strings to None.
    """
    fh = StringIO(csv_text)
    # Detect delimiter (prefer comma; but try to infer)
    sample = fh.read(2048)
    fh.seek(0)
    dialect = None
    try:
        sniffer = csv.Sniffer()
        dialect = sniffer.sniff(sample, delimiters=",;\t")
    except Exception:
        dialect = csv.get_dialect("excel")
    reader = csv.DictReader(fh, dialect=dialect)
    rows: List[Dict[str, Any]] = []
    for raw_row in reader:
        row = {}
        for k, v in raw_row.items():
            if k is None:
                continue
            key = k.strip().lower().replace(" ", "_")
            if isinstance(v, str):
                v2 = v.strip()
                if v2 == "" or v2.lower() == "nan":
                    row[key] = None
                else:
                    row[key] = v2
            else:
                row[key] = v
        # skip completely-empty rows
        if any(v is not None and v != "" for v in row.values()):
            rows.append(row)
    return rows


# -----------------------
# convert-or-ingest (preview)
# -----------------------
@app.post("/api/convert-or-ingest", response_class=PlainTextResponse)
async def convert_or_ingest(file: UploadFile = File(...)):
    filename = file.filename or "uploaded"
    content_type = file.content_type or ""
    content = await file.read()
    size = len(content)
    if size == 0:
        raise HTTPException(status_code=400, detail="Empty file uploaded.")
    if size > MAX_DOCX_SIZE:
        raise HTTPException(status_code=413, detail=f"File too large. Max {MAX_DOCX_SIZE} bytes.")
    lower = filename.lower()

    if lower.endswith(".csv") or "text/csv" in content_type:
        try:
            csv_text = content.decode("utf-8")
        except UnicodeDecodeError:
            csv_text = content.decode("latin-1")
        if "," not in csv_text and ";" not in csv_text and "\t" not in csv_text:
            raise HTTPException(status_code=400, detail="Uploaded file doesn't appear to be a CSV.")
        return PlainTextResponse(content=csv_text, media_type="text/csv", headers={
            "Content-Disposition": f'attachment; filename="{filename if filename.endswith(".csv") else "data.csv"}"'
        })

    if lower.endswith(".docx") or content_type in ALLOWED_DOCX_CONTENT_TYPES:
        try:
            bio = BytesIO(content)
            csv_text = docx_to_csv_text_with_fallback(bio, table_selection="first")
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Conversion failed: {e}")
        return PlainTextResponse(content=csv_text, media_type="text/csv", headers={
            "Content-Disposition": f'attachment; filename="{filename.rsplit(".",1)[0]}.csv"'
        })

    try:
        text = content.decode("utf-8")
        if "," in text or ";" in text or "\t" in text:
            return PlainTextResponse(content=text, media_type="text/csv")
    except Exception:
        pass

    raise HTTPException(status_code=415, detail=f"Unsupported file type: {content_type} / {filename}")


# -----------------------
# Upload endpoint (stage all rows into staging_raw)
# -----------------------
@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...), dataset: str = Form(...)):
    """
    Upload endpoint that STAGES ALL ROWS into staging_raw immediately.
    - Accepts .csv or .docx (docx converts first table -> CSV or paragraphs fallback)
    - Parses CSV into rows (dict per row) and inserts each as a staging_raw row with upload_id
    - Does NOT call ETL cleaning or RPCs here (explicit /api/process should be used)
    """
    dataset_key = dataset.lower().strip()
    if dataset_key not in DATASET_MAP:
        raise HTTPException(status_code=400, detail=f"Unsupported dataset: {dataset}")

    filename = file.filename or "uploaded"
    content_type = file.content_type or ""

    # read uploaded bytes
    content = await file.read()
    size = len(content)
    if size == 0:
        raise HTTPException(status_code=400, detail="Empty file uploaded.")
    if size > MAX_FILE_BYTES:
        raise HTTPException(status_code=413, detail=f"File too large. Max {MAX_FILE_BYTES} bytes.")

    # save original upload to a temporary file (keeps a copy)
    suffix = "." + (filename.split(".")[-1] if "." in filename else "tmp")
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(content)
            tmp.flush()
            tmp_path = tmp.name
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Could not save uploaded file: {e}")

    # create an etl_runs row immediately and get its integer id
    run_id = insert_etl_run(f"upload_{dataset_key}", "staged", note=filename)

    try:
        lower = filename.lower()
        csv_text: Optional[str] = None

        # CSV by extension or content
        if lower.endswith(".csv") or "text/csv" in content_type:
            try:
                csv_text = content.decode("utf-8")
            except UnicodeDecodeError:
                csv_text = content.decode("latin-1")
        elif lower.endswith(".docx") or content_type in ALLOWED_DOCX_CONTENT_TYPES:
            # convert docx to CSV (table first, fallback to paragraphs)
            bio = BytesIO(content)
            csv_text = docx_to_csv_text_with_fallback(bio, table_selection="first")
        else:
            # fallback: try decode and treat as CSV if looks like CSV
            try:
                text = content.decode("utf-8")
            except UnicodeDecodeError:
                text = content.decode("latin-1")
            if "," in text or ";" in text or "\t" in text:
                csv_text = text
            else:
                # unsupported type -> cleanup and error
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass
                raise HTTPException(status_code=415, detail=f"Unsupported file type: {content_type} / {filename}")

        if not csv_text:
            try:
                os.remove(tmp_path)
            except Exception:
                pass
            raise HTTPException(status_code=400, detail="Could not obtain CSV text from upload.")

        # parse CSV into list of dicts
        parsed_rows = parse_csv_text_to_dicts(csv_text)

        # if no rows found, still create a staging_raw pointing to file (so UI can show file)
        if not parsed_rows:
            # insert single staging row pointing to file (raw metadata)
            res = sb.table("staging_raw").insert({
                "entity": dataset_key,
                "raw": {"filename": filename, "note": "no rows parsed"},
                "processed": False,
                "upload_id": run_id,
                "original_filename": filename,
                "file_pointer": tmp_path,
                "detected_entity": dataset_key,
                "notes": {"staged_at": datetime.now(timezone.utc).isoformat()}
            }).execute()
            staged_count = 1 if not (isinstance(res, dict) and res.get("error")) else 0
            safe_update_etl_run(run_id, "staged", note=f"staged_rows={staged_count}")
            return JSONResponse({
                "status": "ok",
                "dataset": dataset_key,
                "filename": filename,
                "upload_id": run_id,
                "staged_rows": staged_count,
                "message": "no rows parsed; staging single file pointer row"
            })

        # --- VALIDATION: compute validation errors for parsed rows (but DO NOT remove them from staging) ---
        valid_rows: List[Dict[str, Any]] = []
        invalid_rows: List[Dict[str, Any]] = []
        validation_map: List[Optional[str]] = []  # parallel to parsed_rows: None or error string
        for r in parsed_rows:
            err = validate_required_fields(dataset_key, r)
            validation_map.append(err)
            if err:
                bad = dict(r)
                bad["_upload_validation_error"] = err
                invalid_rows.append(bad)
            else:
                valid_rows.append(r)

        # Insert invalid rows into import_errors (one row per bad record)
        error_count = 0
        for bad_row in invalid_rows:
            try:
                sb.table("import_errors").insert({
                    "sourcetable": "staging_raw",
                    "sourceid": None,
                    "raw": bad_row if isinstance(bad_row, dict) else {"raw": bad_row},
                    "errormessage": bad_row.get("_upload_validation_error", "validation failed"),
                    "createdat": datetime.now(timezone.utc).isoformat()
                }).execute()
                error_count += 1
            except Exception as _e:
                # log but do not fail the whole upload
                print("Warning: could not insert import_errors row:", str(_e))
                continue

        # --- NEW BEHAVIOR: insert ALL parsed_rows into staging_raw (so UI shows them all),
        # and attach validation error metadata to staged rows that failed validation.
        staging_records = []
        for i, (r, err) in enumerate(zip(parsed_rows, validation_map)):
            notes = {"staged_at": datetime.now(timezone.utc).isoformat()}
            if err:
                # attach the upload validation error into the notes so you can see it in staging_raw
                notes["_upload_validation_error"] = err
                notes["_upload_valid"] = False
            else:
                notes["_upload_valid"] = True

            staging_records.append({
                "entity": dataset_key,
                "raw": r,
                "processed": False,
                "upload_id": run_id,
                "original_filename": filename,
                # keep file pointer on the first row for debugging/reference; subsequent rows set None
                "file_pointer": tmp_path if i == 0 else None,
                "detected_entity": dataset_key,
                "notes": notes
            })

        # insert in batches using batch_insert helper
        staged_count = batch_insert("staging_raw", staging_records)

        # Optionally store original upload to storage (keeps an external copy)
        if STORE_UPLOADS:
            try:
                dest_path = f"uploads/{int(time.time())}_{os.path.basename(filename)}"
                with open(tmp_path, "rb") as fh:
                    data_bytes = fh.read()
                upload_res = sb.storage.from_("uploads").upload(dest_path, data_bytes, {"cacheControl": "3600"})
                if isinstance(upload_res, dict) and upload_res.get("error"):
                    print("Warning: storage upload error:", upload_res.get("error"))
                else:
                    print("Saved original upload to storage:", dest_path)
            except Exception as e:
                print("Warning: storing original upload failed:", str(e))

        # update etl_runs row to staged + note counts (include error_rows)
        safe_update_etl_run(run_id, "staged", note=f"staged_rows={staged_count} error_rows={error_count}")

        return JSONResponse(
            {
                "status": "ok",
                "dataset": dataset_key,
                "filename": filename,
                "upload_id": run_id,
                "staged_rows": staged_count,
                "error_rows": error_count,
                "file_pointer": tmp_path,
            }
        )

    except Exception as e:
        try:
            sb.table("import_errors").insert(
                {
                    "sourcetable": "staging_raw",
                    "sourceid": None,
                    "raw": {"filename": filename, "error": str(e)},
                    "errormessage": "Upload staging exception",
                    "createdat": datetime.now(timezone.utc).isoformat(),
                }
            ).execute()
        except Exception:
            pass
        try:
            safe_update_etl_run(run_id, "failed", note=str(e))
        except Exception:
            pass
        # cleanup temp file on failure
        try:
            os.remove(tmp_path)
        except Exception:
            pass

        # IMPORTANT: always return valid JSON to the frontend (prevents frontend JSON parse errors)
        return JSONResponse(
            status_code=500,
            content={
                "status": "error",
                "message": "Upload staging failed",
                "detail": str(e),
            },
        )

# ---- ALIAS ROUTE: accept upload at /upload as well as /api/upload ----
# This wrapper keeps your existing upload logic identical and only adds a second URL
@app.post("/upload")
async def upload_file_alias(file: UploadFile = File(...), dataset: str = Form(...)):
    """
    Alias for /api/upload to accomodate frontends calling /upload (prevents 404).
    Delegates to the existing upload_file handler.
    """
    return await upload_file(file=file, dataset=dataset)


# -----------------------
# Process endpoint (explicit): process a staged file into cleaned tables + call RPC
# (unchanged from earlier design, except for a tiny alliance normalization right before inserting cleaned rows)
# -----------------------
@app.post("/api/process")
async def process_staged(staging_id: Optional[int] = Form(None), upload_id: Optional[int] = Form(None), dataset: Optional[str] = Form(None)):
    """
    Process a staged upload. Provide either `staging_id` (preferred) OR `upload_id`.
    This will:
      - read the staged file (if file_pointer present) or read rows previously staged
      - run the ETL module's clean_file() if needed OR process staged_raw rows
      - insert cleaned rows into cleaned_table
      - call RPC to promote into dims
    """
    # locate a staging row (to discover file_pointer, detected_entity, upload_id)
    staging_row = None
    if staging_id is not None:
        q = sb.table("staging_raw").select("*").eq("id", staging_id).limit(1).execute()
        if isinstance(q, dict) and q.get("error"):
            raise HTTPException(status_code=500, detail=str(q.get("error")))
        if hasattr(q, "data") and q.data:
            staging_row = q.data[0]
    elif upload_id is not None:
        q = sb.table("staging_raw").select("*").eq("upload_id", upload_id).order("id", desc=False).limit(1).execute()
        if isinstance(q, dict) and q.get("error"):
            raise HTTPException(status_code=500, detail=str(q.get("error")))
        if hasattr(q, "data") and q.data:
            staging_row = q.data[0]

    if not staging_row:
        raise HTTPException(status_code=404, detail="staging row not found for provided staging_id/upload_id")

    detected_entity = (dataset or staging_row.get("detected_entity") or staging_row.get("entity") or "").lower()
    if detected_entity not in DATASET_MAP:
        raise HTTPException(status_code=400, detail=f"Unsupported dataset/detected entity: {detected_entity}")

    cfg = DATASET_MAP[detected_entity]
    cleaned_table = cfg["cleaned_table"]
    rpc_name = cfg["rpc"]
    etl_module = cfg["etl_module"]

    # create a processing etl_runs row (or reuse upload_id's etl_runs if provided)
    run_id = insert_etl_run(f"process_{detected_entity}", "started", note=f"staging_id={staging_row.get('id')}")
    try:
        # First attempt: if there is a file_pointer and the file exists, prefer running etl_module.clean_file(tmp)
        file_pointer = staging_row.get("file_pointer")
        converted_tmp_path = None
        cleaned_rows: List[Dict[str, Any]] = []
        raw_rows: List[Dict[str, Any]] = []

        if file_pointer and os.path.exists(file_pointer):
            # detect extension
            _, ext = os.path.splitext(file_pointer.lower())
            tmp_path_for_etl = file_pointer
            if ext == ".docx":
                with open(file_pointer, "rb") as fh:
                    bio = BytesIO(fh.read())
                    csv_text = docx_to_csv_text_with_fallback(bio, table_selection="first")
                tmp_csv = tempfile.NamedTemporaryFile(delete=False, suffix=".csv")
                tmp_csv.write(csv_text.encode("utf-8"))
                tmp_csv.flush()
                converted_tmp_path = tmp_csv.name
                tmp_csv.close()
                tmp_path_for_etl = converted_tmp_path

            # call the ETL module clean_file
            cleaned_rows, raw_rows = etl_module.clean_file(tmp_path_for_etl)
        else:
            # If no file pointer (or file missing), use previously staged rows (all rows for upload_id)
            q_all = sb.table("staging_raw").select("*").eq("upload_id", staging_row.get("upload_id")).execute()
            rows_data = []
            if isinstance(q_all, dict) and q_all.get("error"):
                raise RuntimeError(q_all.get("error"))
            if hasattr(q_all, "data") and q_all.data:
                rows_data = q_all.data
            # convert staging rows raw-> raw_rows list and also create cleaned_rows by calling etl.cleaner on a temp CSV if possible
            # best-effort: if raw entries look like dicts, pass them as raw_rows and let process_flights_upload or other ETL runtime handle
            raw_rows = []
            for r in rows_data:
                raw_rows.append(r.get("raw") if isinstance(r.get("raw"), dict) else {"rawjson": r.get("raw")})

            # If ETL module provides a row-based runtime, we'll rely on RPC processing step instead of local clean
            cleaned_rows = []

        # If raw_rows was None or empty, coerce
        if raw_rows is None:
            raw_rows = [{"rawjson": r} for r in cleaned_rows]

        # validate shapes
        if not isinstance(cleaned_rows, list):
            cleaned_rows = []
        if not isinstance(raw_rows, list):
            raw_rows = [raw_rows]

        # Insert cleaned rows (attach upload_id) — only if cleaned_rows present
        cleaned_count = 0
        if cleaned_rows:
            cleaned_records = []
            for r in cleaned_rows:
                rec = dict(r)
                rec.setdefault("rawjson", r.get("rawjson", r))
                rec["upload_id"] = staging_row.get("upload_id") or run_id

                # ---------- NEW: normalize alliance in Python ETL: replace missing/null/empty with 'none' ----------
                # This ensures the cleaned table receives 'none' (string) instead of None/NULL for alliance.
                # It intentionally does not treat alliance as required — it just normalizes missing values.
                try:
                    # if alliance key missing or empty/None -> set to 'none'
                    if "alliance" not in rec or rec.get("alliance") is None or str(rec.get("alliance")).strip() == "":
                        rec["alliance"] = "none"
                    else:
                        # normalize whitespace
                        rec["alliance"] = str(rec.get("alliance")).strip()
                except Exception:
                    rec["alliance"] = "none"
                # -----------------------------------------------------------------------------------------------

                cleaned_records.append(rec)
            cleaned_count = batch_insert(cleaned_table, cleaned_records)

        # If cleaned_rows empty, we still want to run the RPC using upload_id so the server-side RPC can consume staged rows
        processed_count = 0
        processed_count = call_rpc_once(rpc_name, p_upload_id=staging_row.get("upload_id") or run_id)

        # Mark staging rows processed (for this upload_id)
        try:
            sb.table("staging_raw").update({"processed": True}).eq("upload_id", staging_row.get("upload_id")).execute()
        except Exception:
            pass

        safe_update_etl_run(run_id, "success", note=f"cleaned_inserted={cleaned_count} processed_into_dims={processed_count}")

        return JSONResponse({
            "status": "ok",
            "dataset": detected_entity,
            "staging_id": staging_row.get("id"),
            "cleaned_inserted": cleaned_count,
            "processed_into_dims": processed_count,
        })

    except Exception as e:
        try:
            sb.table("import_errors").insert(
                {
                    "sourcetable": cleaned_table,
                    "sourceid": None,
                    "raw": {"staging_id": staging_row.get("id"), "error": str(e)},
                    "errormessage": "Processing exception",
                    "createdat": datetime.now(timezone.utc).isoformat(),
                }
            ).execute()
        except Exception:
            pass
        try:
            safe_update_etl_run(run_id, "failed", note=str(e))
        except Exception:
            pass
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        try:
            if 'converted_tmp_path' in locals() and converted_tmp_path and os.path.exists(converted_tmp_path):
                os.remove(converted_tmp_path)
        except Exception:
            pass

# End of file
