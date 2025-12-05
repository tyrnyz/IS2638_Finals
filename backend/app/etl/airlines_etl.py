# backend/app/etl/airlines_etl.py
"""
airlines_etl.clean_file(path) -> (cleaned_rows, raw_rows)

- Accepts a path to a .csv or .docx file.
- Normalizes column names -> snake_case
- Standardizes to:
    airlinekey
    airlinename
    alliance
    rawjson
- Returns:
    cleaned_rows: list[dict]
    raw_rows: list[dict]
"""

from __future__ import annotations
from pathlib import Path
import io
import pandas as pd
import numpy as np
from typing import List, Tuple, Dict, Any, Optional
from datetime import datetime

try:
    from docx import Document
except Exception:
    Document = None

from ..services.supabase_client import sb

def _normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    cols = (
        df.columns.str.strip()
        .str.replace(r"([a-z0-9])([A-Z])", r"\1_\2", regex=True)
        .str.replace(r"[ \-]+", "_", regex=True)
        .str.lower()
    )
    df.columns = cols
    return df


def _df_to_records(df: pd.DataFrame) -> List[Dict[str, Any]]:
    str_cols = df.select_dtypes(include=["object", "string"]).columns
    for c in str_cols:
        df[c] = df[c].astype(str).str.strip().replace({"nan": None, "None": None})
    df = df.where(pd.notnull(df), None)
    return df.to_dict(orient="records")


def _read_docx_table(path: Path) -> pd.DataFrame:
    if Document is None:
        raise RuntimeError("python-docx not installed.")
    doc = Document(str(path))

    if doc.tables:
        table = doc.tables[0]
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            rows.append(cells)

        if len(rows) >= 2:
            header = rows[0]
            data = rows[1:]
            csv_buf = io.StringIO()
            csv_buf.write(",".join(header) + "\n")
            for r in data:
                escaped = []
                for cell in r:
                    if '"' in cell:
                        cell = cell.replace('"', '""')
                    if "," in cell or '"' in cell:
                        escaped.append(f'"{cell}"')
                    else:
                        escaped.append(cell)
                csv_buf.write(",".join(escaped) + "\n")
            csv_buf.seek(0)
            return pd.read_csv(csv_buf, engine="python")

    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not lines:
        raise RuntimeError("DOCX contained no data")
    csv_buf = io.StringIO("\n".join(lines))
    return pd.read_csv(csv_buf, engine="python")


def _read_csv_file(path: Path) -> pd.DataFrame:
    try:
        return pd.read_csv(path, engine="python")
    except Exception:
        return pd.read_csv(path, engine="python", encoding="latin-1")


def clean_file(path: str) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    """
    Returns (cleaned_rows, raw_rows). Uses no-underscore names:
    - cleaned rows: airlinekey, airlinename, alliance, rawjson
    - raw rows: {"rawjson": {...}}
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(path)

    ext = p.suffix.lower()
    if ext == ".docx":
        df = _read_docx_table(p)
    else:
        df = _read_csv_file(p)

    df = _normalize_columns(df)

    # map common variations
    if "airline" in df.columns and "airline_name" not in df.columns:
        df = df.rename(columns={"airline": "airline_name"})

    # trim string cols
    str_cols = df.select_dtypes(include=["object", "string"]).columns
    for c in str_cols:
        df[c] = df[c].astype(str).str.strip().replace({"nan": None, "None": None})

    # airlinekey normalization (from airline_key -> airlinekey)
    if "airline_key" in df.columns:
        df = df.rename(columns={"airline_key": "airlinekey"})
    if "airlinekey" in df.columns:
        df["airlinekey"] = df["airlinekey"].astype(str).str.strip()
        df["airlinekey"] = df["airlinekey"].replace({"nan": None, "": None, "None": None})
        df["airlinekey"] = df["airlinekey"].where(pd.notnull(df["airlinekey"]), None)
        df.loc[df["airlinekey"].notnull(), "airlinekey"] = df.loc[df["airlinekey"].notnull(), "airlinekey"].str.upper()

    # airlinename normalization (from airline_name -> airlinename)
    if "airline_name" in df.columns:
        df = df.rename(columns={"airline_name": "airlinename"})
    if "airlinename" in df.columns:
        df["airlinename"] = df["airlinename"].astype(str).str.strip()
        df["airlinename"] = df["airlinename"].apply(lambda s: s.title() if s and s.lower() not in ("nan", "none") else None)

    # alliance normalization
    if "alliance" in df.columns:
        df["alliance"] = df["alliance"].replace({np.nan: None, "nan": None, "": None})
        df.loc[df["alliance"].notnull(), "alliance"] = df.loc[df["alliance"].notnull(), "alliance"].astype(str).str.strip()

    df = df.drop_duplicates().reset_index(drop=True)

    records = _df_to_records(df)
    cleaned_rows: List[Dict[str, Any]] = []
    raw_rows: List[Dict[str, Any]] = []

    for r in records:
        raw_rows.append({"rawjson": r})
        cleaned_rows.append(
            {
                "airlinekey": r.get("airlinekey"),
                "airlinename": r.get("airlinename"),
                "alliance": r.get("alliance"),
                "rawjson": r,
            }
        )

    return cleaned_rows, raw_rows


# --- ETL runtime entrypoint used by dispatcher/CLI ---
def _upsert_airline_row(row: Dict[str, Any], upload_id: int) -> None:
    """
    Upsert a single normalized airline row into cleaned_airlines.
    Sets processed = true and upload_id for lineage.
    """
    payload = {
        "airlinekey": row.get("airlinekey"),
        "airlinename": row.get("airlinename"),
        "alliance": row.get("alliance"),
        "rawjson": row.get("rawjson") if "rawjson" in row else row,
        "upload_id": upload_id,
        "processed": True,
        "insertedat": datetime.utcnow().isoformat(),
        "error_count": 0,
        "last_error": None
    }
    # use on_conflict = "airlinekey" to upsert by natural key
    res = supabase.table("cleaned_airlines").upsert(payload, on_conflict="airlinekey").execute()
    if getattr(res, "status_code", None) and res.status_code >= 400:
        raise Exception(res.error)


def process_airlines_upload(upload_id: int, raw: Dict[str, Any], run_id: int = None) -> Dict[str, int]:
    """
    Entrypoint for dispatcher/CLI.
    - upload_id: staging_raw id
    - raw: staging_raw.raw (expected to be {"rows": [...]} if parsed)
    Returns: {"processed": n, "errors": m}
    """
    rows = []
    # accept multiple possible shapes:
    # - raw["rows"] is a list of plain dicts (parsed CSV)
    # - raw may be {"raw_rows": [{"rawjson": ...}, ...]}
    if not raw:
        return {"processed": 0, "errors": 0}

    if isinstance(raw, dict) and "rows" in raw and isinstance(raw["rows"], list):
        rows = raw["rows"]
    elif isinstance(raw, dict) and "raw_rows" in raw and isinstance(raw["raw_rows"], list):
        rows = [r.get("rawjson") if isinstance(r, dict) and "rawjson" in r else r for r in raw["raw_rows"]]
    elif isinstance(raw, list):
        rows = raw
    else:
        # unknown shape: try to find any array-ish fields
        for v in raw.values() if isinstance(raw, dict) else []:
            if isinstance(v, list):
                rows = v
                break

    processed = 0
    errors = 0

    for r in rows:
        try:
            # normalize if row is wrapped as {"rawjson": {...}}
            if isinstance(r, dict) and "rawjson" in r:
                rec = r["rawjson"]
            else:
                rec = r
            # attempt to build normalized shape if necessary
            normalized = {
                "airlinekey": rec.get("airlinekey") or rec.get("iata") or rec.get("icao"),
                "airlinename": rec.get("airlinename") or rec.get("airline_name") or rec.get("name"),
                "alliance": rec.get("alliance"),
                "rawjson": rec
            }
            if not normalized["airlinekey"]:
                # if no natural key, skip and log
                errors += 1
                supabase.table("import_errors").insert({
                    "upload_id": upload_id,
                    "row_data": rec,
                    "message": "missing airline key"
                }).execute()
                continue

            _upsert_airline_row(normalized, upload_id)
            processed += 1
        except Exception as e:
            errors += 1
            supabase.table("import_errors").insert({
                "upload_id": upload_id,
                "row_data": r,
                "message": str(e)
            }).execute()

    return {"processed": processed, "errors": errors}


if __name__ == "__main__":
    import sys
    p = sys.argv[1] if len(sys.argv) > 1 else "sample_airlines.csv"
    cleaned, raw = clean_file(p)
    print("Cleaned rows:", len(cleaned))
    print(cleaned[:3])
