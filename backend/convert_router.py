# backend/convert_router.py
from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import PlainTextResponse
from typing import List
from docx import Document
from io import BytesIO, BufferedReader
import io

router = APIRouter()

# limit to 10 MB for conversion (adjust as you prefer)
MAX_DOCX_SIZE = 10 * 1024 * 1024

ALLOWED_DOCX_CONTENT_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    # some clients may send application/octet-stream for docx, but we do stronger check below
}

def escape_csv_cell(cell_text: str) -> str:
    """Escape according to basic CSV rules: double quotes doubled and wrap if needed."""
    if cell_text is None:
        cell_text = ""
    text = cell_text.replace("\r", " ").replace("\n", " ").strip()
    if '"' in text:
        text = text.replace('"', '""')
    if ("," in text) or ('"' in text) or ("\n" in text) or ("\r" in text):
        return f'"{text}"'
    return text

def docx_to_csv_text_with_paragraph_fallback(file_like: BufferedReader, table_selection: str = "first") -> str:
    """
    Convert DOCX to CSV-like text.
    If tables exist, use first table (or all).
    If no tables, attempt a paragraph fallback (lines).
    """
    doc = Document(file_like)

    # 1) If tables exist, convert them
    if doc.tables:
        csv_lines = []
        tables = doc.tables if table_selection == "all" else [doc.tables[0]]
        for t in tables:
            for r in t.rows:
                cells = [c.text for c in r.cells]
                escaped = [escape_csv_cell(c) for c in cells]
                csv_lines.append(",".join(escaped))
            if table_selection == "all":
                csv_lines.append("")  # blank line separates tables
        return "\n".join(csv_lines)

    # 2) Paragraph fallback
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    if not paragraphs:
        raise ValueError("DOCX has no tables and no readable paragraphs to convert to CSV.")
    # if any obvious delimiter exists, return paragraphs joined; else do whitespace split
    text_blob = "\n".join(paragraphs)
    if any(sep in text_blob for sep in [",", ";", "\t"]):
        return text_blob
    rows = []
    for p in paragraphs:
        cols = p.split()
        escaped = [escape_csv_cell(c) for c in cols]
        rows.append(",".join(escaped))
    return "\n".join(rows)

def looks_like_csv(content: bytes, max_probe: int = 4096) -> bool:
    """Try to decode a portion and heuristically decide whether it's CSV-like."""
    if not content:
        return False
    sample = content[:max_probe]
    try:
        text = sample.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = sample.decode("latin-1")
        except Exception:
            return False
    # require at least one newline and one delimiter (comma, semicolon, or tab)
    return ("\n" in text) and any(d in text for d in [",", ";", "\t"])

@router.post("/api/convert-or-ingest", response_class=PlainTextResponse)
async def convert_or_ingest(file: UploadFile = File(...)):
    """
    Accepts a CSV or DOCX. If CSV -> returns the CSV text back.
    If DOCX -> converts first table (or paragraph fallback) to CSV and returns CSV text.
    Response content-type is text/csv.
    """
    filename = file.filename or "uploaded"
    content_type = file.content_type or ""

    content = await file.read()
    size = len(content)
    if size == 0:
        raise HTTPException(status_code=400, detail="Empty file uploaded.")
    if size > MAX_DOCX_SIZE:
        raise HTTPException(status_code=413, detail=f"File too large. Max {MAX_DOCX_SIZE} bytes.")

    lower = filename.lower()

    # 1) Strong DOCX detection: check for ZIP header + extension or known content-type
    is_zip_like = content.startswith(b'PK')
    docx_ext = lower.endswith(".docx")
    docx_ct = content_type in ALLOWED_DOCX_CONTENT_TYPES
    if (docx_ext or docx_ct) and is_zip_like:
        try:
            bio = BytesIO(content)
            csv_text = docx_to_csv_text_with_paragraph_fallback(bio, table_selection="first")
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"DOCX conversion failed: {e}")
        return PlainTextResponse(content=csv_text, media_type="text/csv", headers={
            "Content-Disposition": f'attachment; filename="{filename.rsplit(".",1)[0]}.csv"'
        })

    # 2) If it looks like CSV (heuristic), return as text
    if looks_like_csv(content):
        try:
            csv_text = content.decode("utf-8")
        except UnicodeDecodeError:
            csv_text = content.decode("latin-1")
        return PlainTextResponse(content=csv_text, media_type="text/csv", headers={
            "Content-Disposition": f'attachment; filename="{filename if filename.endswith('.csv') else 'data.csv'}"'
        })

    # 3) Not recognized
    raise HTTPException(status_code=415, detail=f"Unsupported or unrecognized file type: {content_type} / {filename}")